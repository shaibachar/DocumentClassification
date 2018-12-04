
import pandas as pd
import html5lib
import numpy
import re

from os import listdir
from os.path import isfile, join
from itertools import islice, tee
from datetime import date, datetime
from docx import Document


class LoadDocuments(object):
    """
    This Class will load Documents and map them
    """

    def __init__(self):
        pass

    def time_decorator(self, func):
        def wrapper(*args, **kwargs):
            start = datetime.now()
            result = func(*args, **kwargs)
            print("time:", (datetime.now()-start))
            return result
        return wrapper

    def allFiles(self, filePath):
        onlyfiles = [join(filePath, f) for f in listdir(
            filePath) if isfile(join(filePath, f))]
        return onlyfiles


    def readData(self, filePath):
        """ This method will read word document and split it to words. it will remove all non alphabet signes """
        try:
            res = []
            document = Document(filePath)
            for para in document.paragraphs:
                words = para.text.split()
                for w in words:
                    w = re.sub('[^A-Za-zא-ת]+', '', w)
                    if len(w)>0:
                        res.append(w)
            return res
        except Exception as e:
            print("error on load:", filePath, e)

    def readDataVector(self, filePath, diff=1):
        """ This method will create map of vectors (two words with pre defined distance) """
        words = self.readData(filePath)
        di = dict()
        i = 1
        while i < len(words):
            key = words[i-1]+","+words[i]
            di[key] = di.get(key, 0)+1
            i += diff
        return di

    def readDataDict(self, filePath):
        """This method will count all words and return map of word to count"""
        try:
            di = dict()
            document = Document(filePath)
            for para in document.paragraphs:
                words = para.text.split()
                for w in words:
                    w = re.sub('[^A-Za-zא-ת]+', '', w)
                    di[w] = di.get(w, 0)+1
            return di
        except Exception as e:
            print("error on load:", filePath, e)


#loadData = LoadDocuments()
#res = loadData.allFiles("./resources/documents")
#print(loadData.readDataVector("./resources/documents/hello.docx",1))
